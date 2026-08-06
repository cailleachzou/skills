"""Generate Aurora Tech L28 weekly status report #03 (Chinese docx).

Fills the SBY weekly status report template per agents/weekly-status-report.md.
"""
import os
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

SKILL_DIR = r'c:\Users\59620\.claude\skills\tendo-brand'
TEMPLATE = os.path.join(SKILL_DIR, 'references', 'SBY - 每周项目状态报告 (中文).docx')
OUTPUT_DIR = os.path.join(SKILL_DIR, 'test')
OUTPUT = os.path.join(OUTPUT_DIR, 'Tendo - DES-2026-ATL28 - Aurora Tech L28 - 周报 #03.docx')

# Colors from legend (Table 1)
COLOR_BLUE = '00B0F0'    # 已完成
COLOR_GREEN = '00B050'   # 符合预期进度
COLOR_ORANGE = 'FFC000'  # 轻微延误
COLOR_RED = 'FF0000'     # 严重延误


def get_tcs(row):
    """Return list of unique tc elements in row order (skips merged duplicates)."""
    return row._tr.findall(qn('w:tc'))


def clear_set_tc(tc, text, color_rgb=None):
    """Clear tc content and set new text, preserving first run's rPr formatting."""
    paragraphs = tc.findall(qn('w:p'))
    rPr_template = None
    pPr_template = None
    if paragraphs:
        first_p = paragraphs[0]
        pPr = first_p.find(qn('w:pPr'))
        if pPr is not None:
            pPr_template = deepcopy(pPr)
        first_run = first_p.find(qn('w:r'))
        if first_run is not None:
            rPr = first_run.find(qn('w:rPr'))
            if rPr is not None:
                rPr_template = deepcopy(rPr)
        for p in paragraphs:
            tc.remove(p)

    # New paragraph
    new_p = tc.makeelement(qn('w:p'), {})
    if pPr_template is not None:
        new_p.append(pPr_template)

    # New run with cloned formatting
    new_run = tc.makeelement(qn('w:r'), {})
    if rPr_template is not None:
        new_run.append(deepcopy(rPr_template))

    if color_rgb is not None:
        rPr = new_run.find(qn('w:rPr'))
        if rPr is None:
            rPr = tc.makeelement(qn('w:rPr'), {})
            new_run.insert(0, rPr)
        for c in rPr.findall(qn('w:color')):
            rPr.remove(c)
        color_el = tc.makeelement(qn('w:color'), {qn('w:val'): color_rgb})
        rPr.append(color_el)

    lines = str(text).split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            br = tc.makeelement(qn('w:br'), {})
            new_run.append(br)
        t_el = tc.makeelement(qn('w:t'), {})
        t_el.set(qn('xml:space'), 'preserve')
        t_el.text = line
        new_run.append(t_el)

    new_p.append(new_run)
    tc.append(new_p)


def get_color(progress_str, note):
    """Color per agent rules:
       100%+已完成 -> blue; 50-99% on track -> green;
       1-49% or slight delay -> orange; 0%+major issue -> red."""
    p = int(str(progress_str).rstrip('%'))
    if p == 100 and note == '已完成':
        return COLOR_BLUE
    if note == '轻微延误' or (1 <= p <= 49):
        return COLOR_ORANGE
    if 50 <= p <= 99 and note in ('进行中', '符合预期', '符合预期进度'):
        return COLOR_GREEN
    if p == 0 and note == '严重延误':
        return COLOR_RED
    return None


def main():
    doc = Document(TEMPLATE)

    # 1. Title paragraph: #06 -> #03
    title_p = doc.paragraphs[0]
    full_text = ''.join(run.text for run in title_p.runs)
    new_text = full_text.replace('#06', '#03')
    if title_p.runs:
        title_p.runs[0].text = new_text
        for run in title_p.runs[1:]:
            run.text = ''

    # 2. Header table (Table 0): label | value
    t0 = doc.tables[0]
    header_values = [
        'Aurora Tech 上海新办公室装修项目',
        'Aurora Tech 上海科技有限公司',
        '上海市浦东新区世纪大道100号环球金融中心28层',
        '2026年7月30日',
        'Cailleach Zou',
    ]
    for ri, val in enumerate(header_values):
        tcs = get_tcs(t0.rows[ri])
        clear_set_tc(tcs[1], val)  # value cell (tcs[0] = label, preserve)

    # 3. Work progress table (Table 2)
    t2 = doc.tables[2]
    section_rows = {1, 11, 19, 24, 29, 34}  # A. B. C.1. C.2. D. E. titles
    work_items = {
        2:  ('1', 'L28 开放办公区网线敷设', '28/07/26', '30/07/26', '80%', '进行中'),
        3:  ('2', '会议室面板端接',         '29/07/26', '31/07/26', '50%', '轻微延误'),
        12: ('1', '机柜安装与理线',         '30/07/26', '01/08/26', '20%', '进行中'),
    }
    for ri in range(2, len(t2.rows)):
        if ri in section_rows:
            continue
        tcs = get_tcs(t2.rows[ri])
        if len(tcs) < 6:
            continue  # not a standard work-item row
        if ri in work_items:
            no, name, start, end, prog, note = work_items[ri]
            color = get_color(prog, note)
            clear_set_tc(tcs[0], no)
            clear_set_tc(tcs[1], name)
            clear_set_tc(tcs[2], start)
            clear_set_tc(tcs[3], end)
            clear_set_tc(tcs[4], prog)
            clear_set_tc(tcs[5], note, color)
        else:
            for tc in tcs:
                clear_set_tc(tc, '')

    # 4. 主要成果 / 下一步关键计划 (Table 3): 2 cols
    t3 = doc.tables[3]
    achievements = [
        ('完成 L28 开放办公区 120 个信息点网线敷设', '8/2 前完成会议室全部面板端接'),
        ('机柜到位并完成初步安装',                 '8/5 前完成机房设备上架'),
    ]
    for ri in range(1, len(t3.rows)):
        tcs = get_tcs(t3.rows[ri])
        idx = ri - 1
        if idx < len(achievements):
            left, right = achievements[idx]
            clear_set_tc(tcs[0], left)
            clear_set_tc(tcs[1], right)
        else:
            for tc in tcs:
                clear_set_tc(tc, '')

    # 5. 风险与问题 (Table 4): 7 tcs per data row
    t4 = doc.tables[4]
    risk_desc = ('会议室桥架走向与原设计不符。\n\n'
                 '影响：面板端接进度延误 1 天。\n\n'
                 '解决方案：与甲方确认调整走向，8/1 前出变更方案。')
    risk_data = ('1', '风险', risk_desc, '中', '30/07/26', '01/08/26', '')
    for ri in range(1, len(t4.rows)):
        tcs = get_tcs(t4.rows[ri])
        if ri == 1:
            for i, val in enumerate(risk_data):
                if i < len(tcs):
                    clear_set_tc(tcs[i], val)
        else:
            for tc in tcs:
                clear_set_tc(tc, '')

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Saved: {OUTPUT}')


if __name__ == '__main__':
    main()
