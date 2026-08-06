"""Verify Aurora Tech L28 weekly status report #03."""
import os
from docx import Document
from docx.oxml.ns import qn

OUTPUT = r'c:\Users\59620\.claude\skills\tendo-brand\test\Tendo - DES-2026-ATL28 - Aurora Tech L28 - 周报 #03.docx'

COLOR_BLUE = '00B0F0'
COLOR_GREEN = '00B050'
COLOR_ORANGE = 'FFC000'
COLOR_RED = 'FF0000'

results = []


def check(name, ok, actual, expected):
    results.append((name, ok, actual, expected))


def cell_text(tc):
    return ''.join(t.text or '' for t in tc.findall('.//' + qn('w:t')))


def cell_color(tc):
    """Get color of first run with explicit color."""
    for r in tc.findall('.//' + qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            col = rPr.find(qn('w:color'))
            if col is not None:
                return col.get(qn('w:val'))
    return None


def get_tcs(row):
    return row._tr.findall(qn('w:tc'))


def main():
    doc = Document(OUTPUT)

    # ---- Title ----
    title_text = doc.paragraphs[0].text
    check('Title #03', '#03' in title_text and '#06' not in title_text,
          title_text, 'contains #03')

    # ---- Table 0: Header ----
    t0 = doc.tables[0]
    header_expected = [
        ('项目名称', 'Aurora Tech 上海新办公室装修项目'),
        ('客户名称', 'Aurora Tech 上海科技有限公司'),
        ('地点 地址', '上海市浦东新区世纪大道100号环球金融中心28层'),
        ('报告日期', '2026年7月30日'),
        ('项目经理', 'Cailleach Zou'),
    ]
    for ri, (label, val) in enumerate(header_expected):
        tcs = get_tcs(t0.rows[ri])
        actual_label = cell_text(tcs[0]).rstrip(':：').strip()
        actual_val = cell_text(tcs[1])
        check(f'Header R{ri} label', label in actual_label or actual_label in label,
              actual_label, label)
        check(f'Header R{ri} value', actual_val == val, actual_val, val)

    # ---- Table 1: Legend (structure preserved) ----
    t1 = doc.tables[1]
    legend_texts = [cell_text(get_tcs(r)[ci]) for r in t1.rows for ci in range(len(get_tcs(r)))]
    legend_joined = ' '.join(legend_texts)
    check('Legend 蓝色=已完成', '蓝色' in legend_joined and '已完成' in legend_joined,
          legend_joined, '蓝色=已完成')
    check('Legend 绿色=符合预期', '绿色' in legend_joined and '符合预期' in legend_joined,
          legend_joined, '绿色=符合预期')
    check('Legend 橙色=轻微延误', '橙色' in legend_joined and '轻微延误' in legend_joined,
          legend_joined, '橙色=轻微延误')
    check('Legend 红色=严重延误', '红色' in legend_joined and '严重延误' in legend_joined,
          legend_joined, '红色=严重延误')

    # ---- Table 2: Work progress ----
    t2 = doc.tables[2]

    # Section titles preserved
    section_checks = {
        1:  'A.',
        11: 'B.',
        19: 'C.1.',
        24: 'C.2.',
        29: 'D.',
        34: 'E.',
    }
    for ri, label in section_checks.items():
        tcs = get_tcs(t2.rows[ri])
        txt = cell_text(tcs[0])
        check(f'Section row {ri}', txt == label, txt, label)

    # A.1 work item (row 2): L28 开放办公区网线敷设, 80%, 进行中, GREEN
    tcs = get_tcs(t2.rows[2])
    check('A.1 编号', cell_text(tcs[0]) == '1', cell_text(tcs[0]), '1')
    check('A.1 工作项', cell_text(tcs[1]) == 'L28 开放办公区网线敷设',
          cell_text(tcs[1]), 'L28 开放办公区网线敷设')
    check('A.1 开始日期', cell_text(tcs[2]) == '28/07/26', cell_text(tcs[2]), '28/07/26')
    check('A.1 结束日期', cell_text(tcs[3]) == '30/07/26', cell_text(tcs[3]), '30/07/26')
    check('A.1 完成度', cell_text(tcs[4]) == '80%', cell_text(tcs[4]), '80%')
    check('A.1 备注', cell_text(tcs[5]) == '进行中', cell_text(tcs[5]), '进行中')
    check('A.1 颜色(绿)', cell_color(tcs[5]) == COLOR_GREEN,
          cell_color(tcs[5]), COLOR_GREEN)

    # A.2 work item (row 3): 会议室面板端接, 50%, 轻微延误, ORANGE
    tcs = get_tcs(t2.rows[3])
    check('A.2 编号', cell_text(tcs[0]) == '2', cell_text(tcs[0]), '2')
    check('A.2 工作项', cell_text(tcs[1]) == '会议室面板端接',
          cell_text(tcs[1]), '会议室面板端接')
    check('A.2 开始日期', cell_text(tcs[2]) == '29/07/26', cell_text(tcs[2]), '29/07/26')
    check('A.2 结束日期', cell_text(tcs[3]) == '31/07/26', cell_text(tcs[3]), '31/07/26')
    check('A.2 完成度', cell_text(tcs[4]) == '50%', cell_text(tcs[4]), '50%')
    check('A.2 备注', cell_text(tcs[5]) == '轻微延误', cell_text(tcs[5]), '轻微延误')
    check('A.2 颜色(橙)', cell_color(tcs[5]) == COLOR_ORANGE,
          cell_color(tcs[5]), COLOR_ORANGE)

    # B.1 work item (row 12): 机柜安装与理线, 20%, 进行中
    tcs = get_tcs(t2.rows[12])
    check('B.1 编号', cell_text(tcs[0]) == '1', cell_text(tcs[0]), '1')
    check('B.1 工作项', cell_text(tcs[1]) == '机柜安装与理线',
          cell_text(tcs[1]), '机柜安装与理线')
    check('B.1 开始日期', cell_text(tcs[2]) == '30/07/26', cell_text(tcs[2]), '30/07/26')
    check('B.1 结束日期', cell_text(tcs[3]) == '01/08/26', cell_text(tcs[3]), '01/08/26')
    check('B.1 完成度', cell_text(tcs[4]) == '20%', cell_text(tcs[4]), '20%')
    check('B.1 备注', cell_text(tcs[5]) == '进行中', cell_text(tcs[5]), '进行中')
    check('B.1 颜色(橙,1-49%)', cell_color(tcs[5]) == COLOR_ORANGE,
          cell_color(tcs[5]), COLOR_ORANGE)

    # Other A rows (4-10) cleared
    for ri in range(4, 11):
        tcs = get_tcs(t2.rows[ri])
        all_empty = all(cell_text(tc) == '' for tc in tcs)
        check(f'A.{ri-1} cleared', all_empty, [cell_text(tc) for tc in tcs], 'all empty')

    # ---- Table 3: 主要成果 / 下一步计划 ----
    t3 = doc.tables[3]
    # Header preserved
    tcs = get_tcs(t3.rows[0])
    check('T3 header left', '主要成果' in cell_text(tcs[0]), cell_text(tcs[0]), '主要成果')
    check('T3 header right', '下一步' in cell_text(tcs[1]), cell_text(tcs[1]), '下一步关键计划活动')
    # Row 1
    tcs = get_tcs(t3.rows[1])
    check('T3 R1 left', cell_text(tcs[0]) == '完成 L28 开放办公区 120 个信息点网线敷设',
          cell_text(tcs[0]), '完成 L28 开放办公区 120 个信息点网线敷设')
    check('T3 R1 right', cell_text(tcs[1]) == '8/2 前完成会议室全部面板端接',
          cell_text(tcs[1]), '8/2 前完成会议室全部面板端接')
    # Row 2
    tcs = get_tcs(t3.rows[2])
    check('T3 R2 left', cell_text(tcs[0]) == '机柜到位并完成初步安装',
          cell_text(tcs[0]), '机柜到位并完成初步安装')
    check('T3 R2 right', cell_text(tcs[1]) == '8/5 前完成机房设备上架',
          cell_text(tcs[1]), '8/5 前完成机房设备上架')

    # ---- Table 4: 风险与问题 ----
    t4 = doc.tables[4]
    # Header preserved (7 tcs)
    tcs = get_tcs(t4.rows[0])
    check('T4 header 项目', '项目' in cell_text(tcs[0]), cell_text(tcs[0]), '项目')
    check('T4 header 风险描述', '风险描述' in cell_text(tcs[1]), cell_text(tcs[1]), '风险描述')
    check('T4 header 描述影响', '描述' in cell_text(tcs[2]) and '影响' in cell_text(tcs[2]),
          cell_text(tcs[2]), '描述, 影响与解决方案')
    check('T4 header 优先级', '优先级' in cell_text(tcs[3]), cell_text(tcs[3]), '优先级')
    check('T4 header 记录日期', '记录日期' in cell_text(tcs[4]), cell_text(tcs[4]), '记录日期')
    check('T4 header 预计完成', '预计完成' in cell_text(tcs[5]), cell_text(tcs[5]), '预计完成日期')
    check('T4 header 备注', '备注' in cell_text(tcs[6]), cell_text(tcs[6]), '备注')
    # Row 1 data
    tcs = get_tcs(t4.rows[1])
    check('T4 R1 项目', cell_text(tcs[0]) == '1', cell_text(tcs[0]), '1')
    check('T4 R1 风险类型', cell_text(tcs[1]) == '风险', cell_text(tcs[1]), '风险')
    desc = cell_text(tcs[2])
    check('T4 R1 描述含桥架', '桥架走向' in desc, desc, 'contains 桥架走向')
    check('T4 R1 含影响', '延误 1 天' in desc, desc, 'contains 延误 1 天')
    check('T4 R1 含解决方案', '变更方案' in desc, desc, 'contains 变更方案')
    check('T4 R1 优先级', cell_text(tcs[3]) == '中', cell_text(tcs[3]), '中')
    check('T4 R1 记录日期', cell_text(tcs[4]) == '30/07/26', cell_text(tcs[4]), '30/07/26')
    check('T4 R1 预计完成', cell_text(tcs[5]) == '01/08/26', cell_text(tcs[5]), '01/08/26')
    # Row 2+ cleared
    tcs2 = get_tcs(t4.rows[2])
    check('T4 R2 cleared', all(cell_text(tc) == '' for tc in tcs2),
          [cell_text(tc) for tc in tcs2], 'all empty')

    # ---- Table 5: 现场照片 preserved ----
    t5 = doc.tables[5]
    check('T5 现场照片 preserved', len(t5.rows) == 9 and len(get_tcs(t5.rows[0])) == 2,
          (len(t5.rows), len(get_tcs(t5.rows[0]))), '(9, 2)')

    # ---- Section headings (paragraphs) ----
    para_texts = [p.text for p in doc.paragraphs]
    check('Para 工作进展', '工作进展' in para_texts, para_texts, 'contains 工作进展')
    check('Para 主要成果', any('主要成果' in t for t in para_texts), para_texts, 'contains 主要成果')
    check('Para 风险与问题', '风险与问题' in para_texts, para_texts, 'contains 风险与问题')
    check('Para 现场照片', '现场照片' in para_texts, para_texts, 'contains 现场照片')

    # ---- Total tables count (6) ----
    check('Tables count', len(doc.tables) == 6, len(doc.tables), 6)

    # Print results
    print('\n=== Verification ===')
    all_pass = True
    for name, ok, actual, expected in results:
        status = 'PASS' if ok else 'FAIL'
        if not ok:
            all_pass = False
        print(f'{status} | {name} | actual={actual!r} | expected={expected!r}')

    pass_n = sum(1 for _, ok, _, _ in results if ok)
    fail_n = sum(1 for _, ok, _, _ in results if not ok)
    print(f'\nTotal: {len(results)} | Pass: {pass_n} | Fail: {fail_n}')
    print(f'Overall: {"PASS" if all_pass else "FAIL"}')


if __name__ == '__main__':
    main()
